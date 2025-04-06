import PyPDF2
from docx import Document
from typing import Union
import os
import re
import unicodedata

def preprocess_text(text: str, 
                   remove_urls: bool = True,
                   remove_emails: bool = True,
                   remove_special_chars: bool = True,
                   normalize_whitespace: bool = True,
                   normalize_unicode: bool = True,
                   preserve_line_breaks: bool = True) -> str:
    """
    Preprocess text for RAG system storage.
    
    Args:
        text (str): Input text to preprocess
        remove_urls (bool): Whether to remove URLs
        remove_emails (bool): Whether to remove email addresses
        remove_special_chars (bool): Whether to remove special characters
        normalize_whitespace (bool): Whether to normalize whitespace
        normalize_unicode (bool): Whether to normalize Unicode characters
        preserve_line_breaks (bool): Whether to preserve paragraph breaks
        
    Returns:
        str: Preprocessed text
    """
    if not text:
        return text

    # Normalize Unicode characters
    if normalize_unicode:
        text = unicodedata.normalize('NFKC', text)
    
    # Remove URLs
    if remove_urls:
        text = re.sub(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', 
                     ' ', text)
    
    # Remove email addresses
    if remove_emails:
        text = re.sub(r'[\w\.-]+@[\w\.-]+\.\w+', ' ', text)
    
    # Preserve paragraph breaks if requested
    if preserve_line_breaks:
        # Replace multiple newlines with a special token
        text = re.sub(r'\n\s*\n', ' [BREAK] ', text)
    
    # Remove special characters but keep basic punctuation
    if remove_special_chars:
        text = re.sub(r'[^a-zA-Z0-9\s.,!?-]', ' ', text)
    
    # Normalize whitespace
    if normalize_whitespace:
        # Replace multiple spaces with single space
        text = re.sub(r'\s+', ' ', text)
        
        if preserve_line_breaks:
            # Restore paragraph breaks
            text = text.replace('[BREAK]', '\n\n')
    
    return text.strip()

def extract_text_from_pdf(pdf_path: str) -> str:
    """
    Extract text from all pages of a PDF file.
    
    Args:
        pdf_path (str): Path to the PDF file
        
    Returns:
        str: Extracted text from the PDF
        
    Raises:
        FileNotFoundError: If the PDF file doesn't exist
        PyPDF2.errors.PdfReadError: If there's an error reading the PDF
    """
    try:
        # Open the PDF file in binary read mode
        with open(pdf_path, 'rb') as file:
            # Create a PDF reader object
            pdf_reader = PyPDF2.PdfReader(file)
            
            # Initialize an empty string to store the text
            text = ""
            
            # Iterate through all pages and extract text
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
                
            return text.strip()
            
    except FileNotFoundError:
        raise FileNotFoundError(f"The file {pdf_path} was not found")
    except Exception as e:
        raise PyPDF2.errors.PdfReadError(f"Error reading PDF file: {str(e)}")

def extract_text_from_docx(docx_path: str, preprocess: bool = True, **preprocess_kwargs) -> str:
    """
    Extract text from a Word document.
    
    Args:
        docx_path (str): Path to the Word document
        preprocess (bool): Whether to preprocess the extracted text
        **preprocess_kwargs: Keyword arguments for preprocess_text function
        
    Returns:
        str: Extracted text from the Word document
    """
    try:
        doc = Document(docx_path)
        text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text.append(" | ".join(row_text))
        
        text = "\n".join(text)
        if preprocess:
            text = preprocess_text(text, **preprocess_kwargs)
        return text
            
    except FileNotFoundError:
        raise FileNotFoundError(f"The file {docx_path} was not found")
    except Exception as e:
        raise ValueError(f"Error reading Word file: {str(e)}")

def extract_text_from_document(file_path: str, preprocess: bool = True, **preprocess_kwargs) -> str:
    """
    Extract text from either a PDF or Word document based on file extension.
    
    Args:
        file_path (str): Path to the document file
        preprocess (bool): Whether to preprocess the extracted text
        **preprocess_kwargs: Keyword arguments for preprocess_text function
        
    Returns:
        str: Extracted text from the document
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"The file {file_path} was not found")
    
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension in ['.docx', '.doc']:
        return extract_text_from_docx(file_path, preprocess, **preprocess_kwargs)
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")

# text=extract_text_from_pdf(r"E:\RFP-Analyzer\NEW_RESUME (7).pdf")
# print(text)
